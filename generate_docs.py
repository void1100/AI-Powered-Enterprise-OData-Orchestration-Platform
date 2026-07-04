"""Generate comprehensive project documentation in Word format."""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
import os

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title page
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('Advanced OData Service Orchestration Architecture', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle = doc.add_paragraph('MCP-Based Enterprise Chatbot with Multi-Service Integration')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Version 2.0 | July 2026\n').font.size = Pt(12)
meta.add_run('Repository: github.com/Karanpr-18/ODATA\n').font.size = Pt(10)
meta.add_run('Classification: Internal Technical Document').font.size = Pt(10)

doc.add_page_break()

# --- Table of Contents placeholder ---
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Executive Summary',
    '2. System Architecture',
    '3. Tech Stack',
    '4. Core Features',
    '5. Query Engine Deep Dive',
    '6. ML & Analytics Pipeline',
    '7. Security & Authentication',
    '8. Infrastructure & Deployment',
    '9. API Reference',
    '10. System Requirements',
    '11. Cost Analysis',
    '12. Architecture Comparison',
    '13. Recommendations',
    '14. Future Roadmap',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# --- 1. Executive Summary ---
doc.add_heading('1. Executive Summary', level=1)
doc.add_paragraph(
    'Advanced OData Service Orchestration Architecture is a self-contained, locally-runnable '
    'implementation of an enterprise-grade chatbot that connects to any OData v4 service, '
    'orchestrates queries through an LLM reasoning engine, and exposes a chat-style frontend '
    'for natural-language interaction with SAP and enterprise data systems.'
)
doc.add_paragraph(
    'The system supports multiple LLM providers (Groq, NVIDIA, Gemini, OpenRouter, OpenAI), '
    'graph-based entity relationship storage via Neo4j, vector memory for conversational context, '
    'and a full admin portal with role-based access control. It is fully Dockerized for '
    'one-command deployment.'
)

# Key metrics table
doc.add_heading('Key Metrics', level=2)
table = doc.add_table(rows=8, cols=2, style='Light Grid Accent 1')
metrics = [
    ('Metric', 'Value'),
    ('Services Registered', '2 (Sopra PO, PP MPE Order Manage)'),
    ('Total Entities', '166 (8 + 158)'),
    ('ML Algorithms', '17 (5 unsupervised + 12 supervised)'),
    ('LLM Providers', '5 (Groq, NVIDIA, Gemini, OpenRouter, Mock)'),
    ('Docker Services', '4 (Frontend, Backend, Neo4j, n8n)'),
    ('API Endpoints', '40+ REST endpoints'),
    ('Monthly Cost (Prototype)', '~$3-5/month'),
]
for i, (k, v) in enumerate(metrics):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v

doc.add_page_break()

# --- 2. System Architecture ---
doc.add_heading('2. System Architecture', level=1)

doc.add_heading('2.1 High-Level Flow', level=2)
doc.add_paragraph(
    'The architecture follows a layered approach with clear separation of concerns:'
)

flow_text = """
User Query (Browser)
      │
      ▼
┌─────────────────────┐
│  nginx Frontend     │  ← Landing page, Chat UI, Admin Portal
│  (port 3000)        │    Static HTML/CSS/JS, no build step
└──────────┬──────────┘
           │ HTTP /api/*
           ▼
┌─────────────────────┐
│  FastAPI Backend    │  ← REST API, Reasoning Engine, Orchestrator
│  (port 8000)        │    Python 3.10, uvicorn
└──────────┬──────────┘
           │
           ▼
┌─────────────────────────────────────────────────────────────┐
│               MCP-Based Orchestration Layer                  │
│                                                              │
│  Query Optimizer → Reasoning Engine → Orchestrator → OData   │
│       │                │                  │                  │
│   Plan Cache     LLM Provider      Service Manager          │
│   ChromaDB RAG   (Groq/NVIDIA)     (SAP CPI, Generic)       │
│                      │                  │                    │
│              ┌───────┘                  │                    │
│              ▼                          ▼                    │
│     Neo4j Graph DB              OData HTTP Client           │
│     (Entity Relations)          (Basic Auth, Pagination)    │
│              │                          │                    │
│              └──────────► ML Engine ◄───┘                    │
│                    (scikit-learn, xgboost)                   │
└─────────────────────────────────────────────────────────────┘
"""
p = doc.add_paragraph(flow_text)
p.style = doc.styles['Normal']
for run in p.runs:
    run.font.name = 'Consolas'
    run.font.size = Pt(8)

doc.add_heading('2.2 Component Interaction', level=2)
doc.add_paragraph(
    'The backend uses a multi-agent architecture with specialized agents for different tasks:'
)
agents = [
    ('Discovery Agent', 'Scans OData metadata, builds entity graph, stores in Neo4j'),
    ('Reasoning Engine', 'Classifies intent, builds query plan, routes to LLM or mock planner'),
    ('Orchestrator', 'Executes query plan, handles pagination, retries, aggregation'),
    ('Query Optimizer', 'Caches plans, classifies intent, optimizes $select clauses'),
    ('Query RAG', 'Retrieves similar past query plans from ChromaDB as few-shot examples'),
]
for name, desc in agents:
    p = doc.add_paragraph()
    p.add_run(f'{name}: ').bold = True
    p.add_run(desc)

doc.add_page_break()

# --- 3. Tech Stack ---
doc.add_heading('3. Tech Stack', level=1)

doc.add_heading('3.1 Backend', level=2)
table = doc.add_table(rows=14, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Component'
table.rows[0].cells[1].text = 'Technology'
table.rows[0].cells[2].text = 'Purpose'
backend_stack = [
    ('Framework', 'FastAPI + uvicorn', 'REST API server, async request handling'),
    ('Orchestration', 'Custom MCP-based', 'Multi-service query planning and execution'),
    ('Graph DB', 'Neo4j', 'Entity relationships, service registry, join history'),
    ('Vector DB', 'ChromaDB', 'Query plan caching, RAG few-shot retrieval'),
    ('Auth DB', 'SQLite', 'User accounts, roles, audit logs'),
    ('LLM Primary', 'Groq (llama-3.3-70b-versatile)', 'Query plan generation, reasoning'),
    ('LLM Secondary', 'NVIDIA (llama-3.3-70b-instruct)', 'Free alternative provider'),
    ('HTTP Client', 'httpx', 'Async OData requests with Basic Auth'),
    ('ML (Unsupervised)', 'scikit-learn', 'K-Means, Anomaly Detection, Correlation'),
    ('ML (Supervised)', 'scikit-learn + xgboost + catboost', '12 classification/regression algorithms'),
    ('Data Processing', 'pandas + numpy', 'Post-fetch aggregation, data cleaning'),
    ('Auth', 'JWT + bcrypt', 'Password hashing, httpOnly cookies, RBAC'),
    ('Container', 'Docker + Docker Compose', 'Full stack deployment'),
]
for i, (comp, tech, purpose) in enumerate(backend_stack):
    table.rows[i+1].cells[0].text = comp
    table.rows[i+1].cells[1].text = tech
    table.rows[i+1].cells[2].text = purpose

doc.add_heading('3.2 Frontend', level=2)
table = doc.add_table(rows=6, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Component'
table.rows[0].cells[1].text = 'Technology'
table.rows[0].cells[2].text = 'Purpose'
frontend_stack = [
    ('Server', 'nginx 1.27', 'Static file serving, API proxy'),
    ('Chat UI', 'HTML/CSS/JavaScript', 'No build step, instant load'),
    ('Charts', 'Chart.js', 'Pie, Bar, Network graph visualization'),
    ('ML UI', 'Vanilla JS', 'Algorithm selection, training, prediction'),
    ('Admin', 'HTML/CSS/JavaScript', 'User management, RBAC, service registry'),
]
for i, (comp, tech, purpose) in enumerate(frontend_stack):
    table.rows[i+1].cells[0].text = comp
    table.rows[i+1].cells[1].text = tech
    table.rows[i+1].cells[2].text = purpose

doc.add_heading('3.3 Infrastructure', level=2)
table = doc.add_table(rows=5, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Service'
table.rows[0].cells[1].text = 'Port'
table.rows[0].cells[2].text = 'Description'
infra = [
    ('Frontend (nginx)', '3000', 'Landing page, Chat app, Admin portal'),
    ('Backend (FastAPI)', '8000', 'REST API, ML engines, OData client'),
    ('Neo4j', '7474/7687', 'Graph database for entity relationships'),
    ('n8n', '5678', 'Workflow automation, chat sharing'),
]
for i, (svc, port, desc) in enumerate(infra):
    table.rows[i+1].cells[0].text = svc
    table.rows[i+1].cells[1].text = port
    table.rows[i+1].cells[2].text = desc

doc.add_page_break()

# --- 4. Core Features ---
doc.add_heading('4. Core Features', level=1)

features = [
    ('Multi-Service OData Integration', [
        'Register any OData v4 service via URL (with Basic Auth support)',
        'SAP CPI non-standard URL pattern: ?service=X&entity=Y&top=N',
        'Auto-recovery from Neo4j when metadata fetch fails',
        'Dynamic entity property metadata passed to LLM',
        'Currently running: Sopra PO Service (8 entities) + PP MPE Order Manage (158 entities)',
    ]),
    ('LLM-Powered Query Generation', [
        '5 provider support: Groq, NVIDIA, Gemini, OpenRouter, Mock',
        'Real-time LLM model switcher in the UI',
        'Intent classification (read, aggregate, compute, comparison)',
        'Smart $select: only requests relevant columns',
        'Query plan caching with MD5 keys (5-min TTL)',
        'RAG: retrieves similar past plans as few-shot examples',
    ]),
    ('Entity Selector + Auto-Join', [
        'Visual entity selector with property labels (Key, ForeignKey, Date, Measure, Status)',
        'Auto-detection of join keys between selected entities',
        'Priority labels: Primary Key, Foreign Key, Fuzzy Match, Confirmed',
        'Multi-entity chain joins (3+ entities)',
        'Neo4j-stored join relationships for future reference',
        'Chat context: selected entities scope chat queries with runtime join detection',
    ]),
    ('Cross-Service Joins', [
        'Union: Stack rows from multiple services',
        'Match: Inner join by common key',
        'Enrichment: Primary + secondary lookup',
    ]),
    ('ML & Analytics (17 Algorithms)', [
        'Unsupervised: Summary, Anomaly, Correlation, K-Means, Feature Importance',
        'Supervised: Decision Tree, Random Forest, XGBoost, CatBoost, SVM, KNN, etc.',
        'Business Insights: automatically generates actionable recommendations',
        'Predictive queries: use trained models in chat queries',
        'Data cleaning pipeline: missing values, outliers, normalization, encoding',
    ]),
    ('LLM + ML Integration', [
        'Data Profiler: scans fetched data for column types, correlations, outliers, distributions, quality score',
        'Auto-Train: after every query with 10+ rows, profiles data, selects best algorithm, trains, then deletes model',
        'Algorithm Selection: Decision Tree (small data), Random Forest (balanced), Gradient Boosting (complex), Linear Regression (normal)',
        'On-Demand Insights: click "Analyze" button for LLM-powered data insights, ML recommendations, chart suggestions',
        'Insights Panel: shows summary, profile stats, correlations, outliers, ML recommendation, clickable suggestion chips',
    ]),
    ('Token Usage Tracker', [
        'Per-query logging: provider, tokens, latency, intent, cached status stored in SQLite',
        'Admin Usage Dashboard: summary cards, provider breakdown bar chart, daily chart, recent queries table',
        'Cost estimation per provider (Groq $0.59/$0.79 per M tokens, NVIDIA free)',
        'API endpoint: GET /admin/usage returns full usage data',
    ]),
    ('Security & Authentication', [
        'JWT tokens with httpOnly cookies',
        'Password strength validation, account lockout after 5 failures',
        'Role-based access control (5 default roles)',
        'Optional auth on chat endpoints',
        'Audit logging for all admin actions',
    ]),
    ('Chat Sharing via n8n', [
        'Share queries via Email, WhatsApp, Slack, or Copy Link',
        'n8n workflow automation for webhook processing',
    ]),
]

for feature_name, items in features:
    doc.add_heading(feature_name, level=2)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# --- 5. Query Engine Deep Dive ---
doc.add_heading('5. Query Engine Deep Dive', level=1)

doc.add_heading('5.1 Query Processing Pipeline', level=2)
steps = [
    'User types natural language query',
    'Query Optimizer classifies intent (read/aggregate/compute/comparison)',
    'Reasoning Engine detects explicit service mentions',
    'If explicit: builds plan directly (0 tokens). If ambiguous: calls LLM',
    'LLM generates query plan with entity, filters, selects',
    'Orchestrator executes plan via OData HTTP client',
    'Post-fetch aggregation: GROUP BY computed in Python',
    'Response sanitizer: filters OData metadata, orders columns',
    'Results returned with chart recommendations and suggestions',
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f'{i}. {step}')

doc.add_heading('5.2 Token Optimization', level=2)
doc.add_paragraph(
    'The system is designed to minimize LLM token consumption per query:'
)
optimizations = [
    'Explicit service detection: mentions "from Sopra PO Service" → 0 LLM tokens',
    'Smart $select: only requests columns relevant to the query',
    'Plan caching: identical queries return cached plans (0 tokens)',
    'RAG: similar past plans reduce prompt complexity',
    'Mock planner: handles simple read queries without LLM',
]
for opt in optimizations:
    doc.add_paragraph(opt, style='List Bullet')

doc.add_heading('5.3 SAP CPI Integration', level=2)
doc.add_paragraph(
    'Special handling for SAP Cloud Platform Integration (CPI) services:'
)
sap_features = [
    'Non-standard URL pattern: ?service=X&entity=Y&top=N',
    'Basic Auth authentication with username/password',
    'Metadata recovery from Neo4j when metadata endpoint returns 500',
    'Auto-cap $top at 2 for entities with strict row limits',
    'Health probe includes auth headers for degraded status detection',
]
for feat in sap_features:
    doc.add_paragraph(feat, style='List Bullet')

doc.add_page_break()

# --- 6. ML & Analytics Pipeline ---
doc.add_heading('6. ML & Analytics Pipeline', level=1)

doc.add_heading('6.1 Algorithm Coverage', level=2)
table = doc.add_table(rows=18, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Algorithm'
table.rows[0].cells[1].text = 'Type'
table.rows[0].cells[2].text = 'Use Case'
ml_algos = [
    ('Summary Statistics', 'Unsupervised', 'Data overview, mean/median/std'),
    ('Anomaly Detection', 'Unsupervised', 'Outlier identification'),
    ('Correlation Analysis', 'Unsupervised', 'Feature relationship discovery'),
    ('K-Means Clustering', 'Unsupervised', 'Data grouping with scatter visualization'),
    ('Feature Importance', 'Unsupervised', 'Key variable identification'),
    ('Decision Tree', 'Supervised', 'Interpretable classification'),
    ('Random Forest', 'Supervised', 'Ensemble classification'),
    ('Linear Regression', 'Supervised', 'Continuous value prediction'),
    ('Logistic Regression', 'Supervised', 'Binary classification (Yes/No)'),
    ('XGBoost', 'Supervised', 'High-performance gradient boosting'),
    ('CatBoost', 'Supervised', 'Categorical feature handling'),
    ('KNN', 'Supervised', 'Instance-based classification'),
    ('SVM', 'Supervised', 'Margin-based classification'),
    ('Gradient Boosting', 'Supervised', 'Sequential ensemble'),
    ('Ada Boost', 'Supervised', 'Weak learner boosting'),
    ('Extra Trees', 'Supervised', 'Randomized tree ensemble'),
    ('Naive Bayes', 'Supervised', 'Probabilistic classification'),
]
for i, (algo, type_, use) in enumerate(ml_algos):
    table.rows[i+1].cells[0].text = algo
    table.rows[i+1].cells[1].text = type_
    table.rows[i+1].cells[2].text = use

doc.add_heading('6.2 Business Insights Engine', level=2)
doc.add_paragraph(
    'The supervised ML results include an automated Business Insights generator that:'
)
insights = [
    'Analyzes model accuracy and feature importance',
    'Generates domain-specific actionable recommendations',
    'Provides risk assessment based on prediction distribution',
    'Suggests next steps for data collection or process improvement',
]
for insight in insights:
    doc.add_paragraph(insight, style='List Bullet')

doc.add_heading('6.3 Predictive Queries', level=2)
doc.add_paragraph(
    'Trained ML models can be used directly in chat queries:'
)
predictive = [
    'Train a model on query results (e.g., classification of purchase orders)',
    'Use the model in subsequent queries: "Predict if PO 4500000001 will be approved"',
    'Supports any service: predictions work across all registered OData services',
    'Results show Yes/No with confidence scores',
]
for pred in predictive:
    doc.add_paragraph(pred, style='List Bullet')

doc.add_page_break()

# --- 7. Security & Authentication ---
doc.add_heading('7. Security & Authentication', level=1)

doc.add_heading('7.1 Authentication System', level=2)
auth_features = [
    'JWT tokens stored in httpOnly cookies (XSS-resistant)',
    'bcrypt password hashing with salt',
    'Password strength validation (min 8 chars, uppercase, lowercase, digit, special)',
    'Account lockout after 5 consecutive failed attempts',
    'Session-based authentication with automatic expiry',
]
for feat in auth_features:
    doc.add_paragraph(feat, style='List Bullet')

doc.add_heading('7.2 Role-Based Access Control', level=2)
table = doc.add_table(rows=6, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Role'
table.rows[0].cells[1].text = 'Permissions'
table.rows[0].cells[2].text = 'Description'
roles = [
    ('Super Admin', 'ALL', 'Full system access'),
    ('Admin', 'manage_users, manage_services, view_analytics', 'User and service management'),
    ('Data Analyst', 'execute_queries, view_services, export_data', 'Query execution and export'),
    ('Viewer', 'view_services', 'Read-only access to services'),
    ('Guest', 'execute_queries', 'Basic query execution only'),
]
for i, (role, perms, desc) in enumerate(roles):
    table.rows[i+1].cells[0].text = role
    table.rows[i+1].cells[1].text = perms
    table.rows[i+1].cells[2].text = desc

doc.add_heading('7.3 Audit Logging', level=2)
doc.add_paragraph(
    'All admin actions are logged with timestamp, user, action, and IP address. '
    'Audit logs are stored in SQLite and viewable in the Admin Portal.'
)

doc.add_page_break()

# --- 8. Infrastructure & Deployment ---
doc.add_heading('8. Infrastructure & Deployment', level=1)

doc.add_heading('8.1 Docker Architecture', level=2)
doc.add_paragraph(
    'The entire stack runs in Docker containers with a single command:'
)
doc.add_paragraph('docker compose up -d --build', style='Intense Quote')

doc.add_heading('8.2 Container Specifications', level=2)
table = doc.add_table(rows=5, cols=4, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Container'
table.rows[0].cells[1].text = 'Image'
table.rows[0].cells[2].text = 'Resource Limit'
table.rows[0].cells[3].text = 'Volume'
containers = [
    ('odata-frontend', 'nginx:1.27-alpine', '256MB RAM', 'Static files'),
    ('odata-backend', 'python:3.10-slim', '2GB RAM', 'backend_data'),
    ('odata-neo4j', 'neo4j:5', '2GB RAM', 'neo4j_data'),
    ('odata-n8n', 'n8nio/n8n', '1GB RAM', 'n8n_data'),
]
for i, (name, img, res, vol) in enumerate(containers):
    table.rows[i+1].cells[0].text = name
    table.rows[i+1].cells[1].text = img
    table.rows[i+1].cells[2].text = res
    table.rows[i+1].cells[3].text = vol

doc.add_heading('8.3 Data Persistence', level=2)
volumes = [
    'neo4j_data: Graph database (entity relationships, service registry)',
    'backend_data: SQLite auth database (users, roles, audit logs)',
    'chroma_cache: ChromaDB vector store (query plan cache, RAG)',
    'n8n_data: n8n workflow configurations',
]
for vol in volumes:
    doc.add_paragraph(vol, style='List Bullet')

doc.add_page_break()

# --- 9. API Reference ---
doc.add_heading('9. API Reference', level=1)

doc.add_heading('9.1 Chat & Analysis', level=2)
table = doc.add_table(rows=9, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Endpoint'
table.rows[0].cells[1].text = 'Method'
table.rows[0].cells[2].text = 'Purpose'
chat_apis = [
    ('/chat', 'POST', 'Natural-language query'),
    ('/analyze', 'POST', 'Unsupervised ML analysis'),
    ('/ml/train', 'POST', 'Train supervised ML model'),
    ('/ml/clean', 'POST', 'Data cleaning pipeline'),
    ('/ml/predict', 'POST', 'Predict using trained model'),
    ('/ml/algorithms', 'GET', 'List supported algorithms'),
    ('/ml/models', 'GET', 'List trained models'),
    ('/share', 'POST', 'Share chat via n8n webhook'),
]
for i, (ep, method, purpose) in enumerate(chat_apis):
    table.rows[i+1].cells[0].text = ep
    table.rows[i+1].cells[1].text = method
    table.rows[i+1].cells[2].text = purpose

doc.add_heading('9.2 Entity Selector & Auto-Join', level=2)
table = doc.add_table(rows=4, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Endpoint'
table.rows[0].cells[1].text = 'Method'
table.rows[0].cells[2].text = 'Purpose'
entity_apis = [
    ('/entities/{service_id}', 'GET', 'Get entities with labeled properties'),
    ('/entities/auto-join', 'POST', 'Detect joins between selected entities'),
    ('/entities/execute-join', 'POST', 'Execute multi-entity join query'),
]
for i, (ep, method, purpose) in enumerate(entity_apis):
    table.rows[i+1].cells[0].text = ep
    table.rows[i+1].cells[1].text = method
    table.rows[i+1].cells[2].text = purpose

doc.add_page_break()

# --- 10. System Requirements ---
doc.add_heading('10. System Requirements', level=1)

doc.add_heading('10.1 Minimum Requirements (Works)', level=2)
table = doc.add_table(rows=8, cols=2, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Component'
table.rows[0].cells[1].text = 'Specification'
min_reqs = [
    ('Operating System', 'Windows 10/11 (64-bit), macOS 12+, or Ubuntu 20.04+'),
    ('CPU', '4 cores / 2 GHz (Intel i5 or AMD Ryzen 5 equivalent)'),
    ('RAM', '8 GB total (6 GB available for Docker)'),
    ('Disk', '10 GB free space (SSD recommended)'),
    ('Docker', 'Docker Desktop 4.0+ or Docker Engine 20.10+'),
    ('Docker Compose', 'v2.0+ (included with Docker Desktop)'),
    ('Network', 'Internet connection (for LLM API calls)'),
]
for i, (comp, spec) in enumerate(min_reqs):
    table.rows[i+1].cells[0].text = comp
    table.rows[i+1].cells[1].text = spec

doc.add_heading('10.2 Recommended Requirements (Smooth)', level=2)
table = doc.add_table(rows=8, cols=2, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Component'
table.rows[0].cells[1].text = 'Specification'
rec_reqs = [
    ('Operating System', 'Windows 11, macOS 13+, or Ubuntu 22.04+'),
    ('CPU', '8 cores / 3 GHz (Intel i7 or AMD Ryzen 7 equivalent)'),
    ('RAM', '16 GB total (12 GB available for Docker)'),
    ('Disk', '25 GB free space (NVMe SSD recommended)'),
    ('Docker', 'Docker Desktop 4.20+'),
    ('Network', '100 Mbps broadband connection'),
    ('Browser', 'Latest Chrome, Firefox, or Edge'),
]
for i, (comp, spec) in enumerate(rec_reqs):
    table.rows[i+1].cells[0].text = comp
    table.rows[i+1].cells[1].text = spec

doc.add_heading('10.3 Enterprise Requirements (Production)', level=2)
table = doc.add_table(rows=8, cols=2, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Component'
table.rows[0].cells[1].text = 'Specification'
ent_reqs = [
    ('Operating System', 'Ubuntu 22.04 LTS or RHEL 8+'),
    ('CPU', '16+ cores / 3 GHz (Intel Xeon or AMD EPYC)'),
    ('RAM', '64 GB ECC RAM'),
    ('Disk', '500 GB NVMe SSD (RAID 10 recommended)'),
    ('Docker', 'Docker Engine 24+ with Swarm or Kubernetes'),
    ('Network', '1 Gbps dedicated, low-latency connection'),
    ('GPU', 'Optional: NVIDIA A10G 24GB (for local embeddings)'),
]
for i, (comp, spec) in enumerate(ent_reqs):
    table.rows[i+1].cells[0].text = comp
    table.rows[i+1].cells[1].text = spec

doc.add_heading('10.4 Actual RAM Consumption (Live Measurements)', level=2)
p = doc.add_paragraph('Measured on development machine: Intel i7-14650HX (16 cores, 24 threads), 31.7 GB RAM, Windows 11')
p.runs[0].font.size = Pt(9)
p.runs[0].font.italic = True

table = doc.add_table(rows=5, cols=5, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Service'
table.rows[0].cells[1].text = 'RAM (Idle)'
table.rows[0].cells[2].text = 'RAM (Under Load)'
table.rows[0].cells[3].text = 'CPU (Idle)'
table.rows[0].cells[4].text = 'CPU (Load)'
ram = [
    ('odata-backend', '890 MB', '~1.2 GB', '0.18%', '10-30%'),
    ('odata-neo4j', '830 MB', '~1.5 GB', '2.81%', '5-15%'),
    ('odata-n8n', '323 MB', '~500 MB', '0.20%', '5-10%'),
    ('odata-frontend', '19 MB', '~50 MB', '0.00%', '1-5%'),
]
for i, (svc, idle, load, cpu_idle, cpu_load) in enumerate(ram):
    table.rows[i+1].cells[0].text = svc
    table.rows[i+1].cells[1].text = idle
    table.rows[i+1].cells[2].text = load
    table.rows[i+1].cells[3].text = cpu_idle
    table.rows[i+1].cells[4].text = cpu_load

p = doc.add_paragraph()
p.add_run('Total: ').bold = True
p.add_run('~2.06 GB idle, ~3.2 GB under load, ~3.2% CPU idle, 20-60% CPU load')

doc.add_heading('10.5 Actual Disk Usage (Live Measurements)', level=2)
table = doc.add_table(rows=7, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Component'
table.rows[0].cells[1].text = 'Size'
table.rows[0].cells[2].text = 'Notes'
disk = [
    ('Project source code', '1.6 MB', 'Backend (1.3 MB) + Frontend (0.3 MB)'),
    ('Docker images', '32.69 GB', '69% reclaimable (22.8 GB)'),
    ('Docker build cache', '21.34 GB', 'All reclaimable'),
    ('Docker containers', '701 MB', '4 running containers'),
    ('Docker volumes', '3.16 GB', 'Neo4j, ChromaDB, n8n, backend data'),
    ('Total Docker', '~57 GB', 'Images + cache + volumes'),
]
for i, (comp, size, desc) in enumerate(disk):
    table.rows[i+1].cells[0].text = comp
    table.rows[i+1].cells[1].text = size
    table.rows[i+1].cells[2].text = desc

doc.add_heading('10.6 What to Delete to Save Space', level=2)
table = doc.add_table(rows=4, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'What to Delete'
table.rows[0].cells[1].text = 'Command'
table.rows[0].cells[2].text = 'Space Recovered'
delete_items = [
    ('Docker build cache', 'docker builder prune -a', '~21 GB'),
    ('Unused Docker images', 'docker image prune -a', '~22 GB'),
    ('Docker volumes (careful!)', 'docker volume prune', '~2.4 GB'),
]
for i, (what, cmd, space) in enumerate(delete_items):
    table.rows[i+1].cells[0].text = what
    table.rows[i+1].cells[1].text = cmd
    table.rows[i+1].cells[2].text = space
p = doc.add_paragraph('Warning: docker volume prune will delete Neo4j data, ChromaDB cache, and n8n workflows. Back up first.')
p.runs[0].font.size = Pt(9)
p.runs[0].font.italic = True

doc.add_heading('10.7 Software Dependencies', level=2)
table = doc.add_table(rows=4, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Software'
table.rows[0].cells[1].text = 'Version'
table.rows[0].cells[2].text = 'Purpose'
sw = [
    ('Docker Desktop', '4.0+', 'Container runtime'),
    ('Docker Compose', 'v2.0+', 'Multi-container orchestration'),
    ('Git', '2.30+', 'Version control'),
]
for i, (soft, ver, purpose) in enumerate(sw):
    table.rows[i+1].cells[0].text = soft
    table.rows[i+1].cells[1].text = ver
    table.rows[i+1].cells[2].text = purpose

doc.add_heading('10.8 Network Requirements', level=2)
table = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Requirement'
table.rows[0].cells[1].text = 'Specification'
net = [
    ('Inbound Ports', '3000 (frontend), 8000 (API), 7474 (Neo4j), 7687 (bolt), 5678 (n8n)'),
    ('Outbound Ports', '443 (HTTPS for LLM APIs)'),
    ('Bandwidth', '10 Mbps minimum for API calls'),
    ('Latency', '<100ms to LLM providers (Groq, NVIDIA, Gemini)'),
]
for i, (req, spec) in enumerate(net):
    table.rows[i+1].cells[0].text = req
    table.rows[i+1].cells[1].text = spec

doc.add_page_break()

# --- 11. Cost Analysis ---
doc.add_heading('11. Cost Analysis', level=1)

doc.add_heading('11.1 LLM API Pricing', level=2)
table = doc.add_table(rows=6, cols=5, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Provider'
table.rows[0].cells[1].text = 'Model'
table.rows[0].cells[2].text = 'Input Cost'
table.rows[0].cells[3].text = 'Output Cost'
table.rows[0].cells[4].text = 'Free Tier'
pricing = [
    ('Groq', 'llama-3.3-70b-versatile', '$0.59/M tokens', '$0.79/M tokens', '14,400 RPD'),
    ('Groq', 'llama-3.1-8b-instant', '$0.05/M tokens', '$0.08/M tokens', '14,400 RPD'),
    ('NVIDIA', 'llama-3.3-70b-instruct', 'Free', 'Free', 'Unlimited'),
    ('Gemini', 'gemini-2.0-flash', '$0.075/M tokens', '$0.30/M tokens', '1500 RPD'),
    ('Gemini', 'gemini-2.5-pro', '$1.25/M tokens', '$10.00/M tokens', '50 RPD'),
]
for i, (prov, model, inp, out, free) in enumerate(pricing):
    table.rows[i+1].cells[0].text = prov
    table.rows[i+1].cells[1].text = model
    table.rows[i+1].cells[2].text = inp
    table.rows[i+1].cells[3].text = out
    table.rows[i+1].cells[4].text = free

doc.add_heading('11.2 Cost Per Query Estimate', level=2)
table = doc.add_table(rows=6, cols=5, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Query Type'
table.rows[0].cells[1].text = 'Tokens (est.)'
table.rows[0].cells[2].text = 'Groq Cost'
table.rows[0].cells[3].text = 'NVIDIA Cost'
table.rows[0].cells[4].text = 'Gemini Flash'
cost_per_query = [
    ('Simple read', '~600', '$0.0004', '$0.00', '$0.00005'),
    ('Aggregation', '~800', '$0.0005', '$0.00', '$0.00006'),
    ('Complex query', '~1200', '$0.0008', '$0.00', '$0.0001'),
    ('Multi-entity join', '~1500', '$0.001', '$0.00', '$0.0001'),
    ('ML training request', '~2000', '$0.0013', '$0.00', '$0.0002'),
]
for i, (qt, tok, groq, nv, gem) in enumerate(cost_per_query):
    table.rows[i+1].cells[0].text = qt
    table.rows[i+1].cells[1].text = tok
    table.rows[i+1].cells[2].text = groq
    table.rows[i+1].cells[3].text = nv
    table.rows[i+1].cells[4].text = gem

doc.add_heading('11.3 Monthly Cost Estimates', level=2)
table = doc.add_table(rows=5, cols=5, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Usage Level'
table.rows[0].cells[1].text = 'Queries/Day'
table.rows[0].cells[2].text = 'Groq Cost'
table.rows[0].cells[3].text = 'NVIDIA Cost'
table.rows[0].cells[4].text = 'Gemini Flash'
monthly_cost = [
    ('Prototype', '10', '~$1', '$0', '~$0.10'),
    ('Development', '100', '~$5', '$0', '~$0.50'),
    ('Small Business', '1,000', '~$50', '$0', '~$5'),
    ('Enterprise', '10,000', '~$374', '$0', '~$37'),
]
for i, (level, qpd, groq, nv, gem) in enumerate(monthly_cost):
    table.rows[i+1].cells[0].text = level
    table.rows[i+1].cells[1].text = qpd
    table.rows[i+1].cells[2].text = groq
    table.rows[i+1].cells[3].text = nv
    table.rows[i+1].cells[4].text = gem

doc.add_heading('11.4 Enterprise Cost Comparison', level=2)
table = doc.add_table(rows=5, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Tier'
table.rows[0].cells[1].text = 'Infrastructure'
table.rows[0].cells[2].text = 'Total Monthly'
ent_cost = [
    ('Prototype (local)', '$0', '~$5'),
    ('Small Enterprise (1K/day)', '~$500', '~$540'),
    ('Full Enterprise (10K/day)', '~$1,530', '~$1,830'),
    ('Full Enterprise (Gemini)', '~$800', '~$1,400'),
]
for i, (tier, infra, total) in enumerate(ent_cost):
    table.rows[i+1].cells[0].text = tier
    table.rows[i+1].cells[1].text = infra
    table.rows[i+1].cells[2].text = total

doc.add_page_break()

# --- 12. Architecture Comparison ---
doc.add_heading('12. Architecture Comparison', level=1)
doc.add_paragraph(
    'Comparison between the current stack (Groq + Neo4j) and alternative approaches:'
)

table = doc.add_table(rows=11, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = 'Dimension'
table.rows[0].cells[1].text = 'Current (Groq + Neo4j)'
table.rows[0].cells[2].text = 'Alternative (Gemini + SurrealDB)'
comparison = [
    ('LLM Quality', 'Llama 3.3 70B (excellent)', 'Gemini 2.5 Pro (best-in-class)'),
    ('LLM Speed', 'Groq is fastest globally', 'Gemini Flash is very fast'),
    ('Data Privacy', 'Embeddings never leave org', 'All text sent to Google'),
    ('Infrastructure', 'Medium (Neo4j container)', 'Low (single SurrealDB)'),
    ('Prototype Cost', '~$5/month', '~$1/month'),
    ('Enterprise Cost', '~$1,830/month', '~$905/month'),
    ('Vendor Lock-in', 'Low (open models)', 'Medium (Google APIs)'),
    ('Context Window', '128K tokens', '1M tokens (Gemini)'),
    ('Multi-modal', 'Text only', 'Images, PDFs, Audio'),
    ('Graph Relationships', 'Native Neo4j Cypher', 'SurrealDB graph edges'),
]
for i, (dim, curr, alt) in enumerate(comparison):
    table.rows[i+1].cells[0].text = dim
    table.rows[i+1].cells[1].text = curr
    table.rows[i+1].cells[2].text = alt

doc.add_page_break()

# --- 13. Recommendations ---
doc.add_heading('13. Recommendations', level=1)

doc.add_heading('13.1 Use Current Stack When:', level=2)
recs_current = [
    'Data sovereignty / intranet deployment is required',
    'Embeddings must not leave the organisation\'s network',
    'Real-time UX speed is the top priority (Groq is ~10x faster)',
    'Offline or air-gapped deployments are needed',
    'Graph relationships need native Cypher queries (Neo4j)',
]
for rec in recs_current:
    doc.add_paragraph(rec, style='List Bullet')

doc.add_heading('13.2 Consider Switching When:', level=2)
recs_switch = [
    'Cost reduction at scale is the priority (51% cheaper with Gemini)',
    'Simplifying infrastructure is desired (eliminate Neo4j + separate DBs)',
    'Better reasoning quality is needed (Gemini 2.5 Pro outperforms Llama 70B)',
    'Multimodal inputs (documents, images from SAP) are required',
    '1M-token context window is beneficial (full OData schema in context)',
]
for rec in recs_switch:
    doc.add_paragraph(rec, style='List Bullet')

doc.add_heading('13.3 Quick Wins for Improvement', level=2)
quick_wins = [
    'Add OpenRouter key rotation for better rate limit handling',
    'Implement session ownership (bind sessions to user IDs)',
    'Enable n8n Email/WhatsApp nodes for automated sharing',
    'Fine-tune LLM on accumulated RAG query plan data',
    'Add caching for OData metadata (reduce startup recovery time)',
    'Implement horizontal scaling with Docker Swarm or Kubernetes',
]
for win in quick_wins:
    doc.add_paragraph(win, style='List Bullet')

doc.add_page_break()

# --- 14. Future Roadmap ---
doc.add_heading('14. Future Roadmap', level=1)

roadmap = [
    ('Phase 1: Production Readiness', [
        'Implement rate limiting and usage metering',
        'Add comprehensive error handling and monitoring',
        'Set up Grafana/Prometheus for observability',
        'Implement backup and disaster recovery for Neo4j',
    ]),
    ('Phase 2: Enterprise Features', [
        'Multi-tenant architecture with data isolation',
        'SSO integration (LDAP, SAML, OAuth)',
        'Advanced audit logging with compliance reporting',
        'API versioning and backward compatibility',
    ]),
    ('Phase 3: Intelligence Enhancement', [
        'Fine-tuned LLM on enterprise-specific OData patterns',
        'Automated anomaly detection and alerting',
        'Natural language to SQL translation for non-OData sources',
        'Predictive analytics with time-series forecasting',
    ]),
    ('Phase 4: Scale & Performance', [
        'Kubernetes deployment with auto-scaling',
        'CDN for frontend static assets',
        'Redis caching layer for hot queries',
        'Async processing with message queues (RabbitMQ)',
    ]),
]

for phase_name, items in roadmap:
    doc.add_heading(phase_name, level=2)
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

# Save
output_path = r'C:\Users\Lenovo\OneDrive\Desktop\projects\Advance Odata service orchestrated architecture 2\Advanced_OData_Architecture_Documentation.docx'
doc.save(output_path)
print(f"Document saved to: {output_path}")
